"""DOCX / DOC 文件解析器"""

from __future__ import annotations

import logging
import re
import struct
import subprocess
import tempfile
from pathlib import Path

from docx import Document as DocxDocument
from docx.oxml.ns import qn

from ..models import Document, ContentType
from ..config import Config
from .base import BaseParser

logger = logging.getLogger("omniparser.parsers.docx")


class DocxParser(BaseParser):
    """解析 .docx 和 .doc 文件

    - .docx 直接用 python-docx 解析
    - .doc 优先用 libreoffice 转换，不可用时降级为 olefile 纯 Python 提取
    """

    supported_extensions = {".docx", ".doc"}

    def parse(self, file_path: Path) -> list[Document]:
        if file_path.suffix.lower() == ".doc":
            return self._parse_doc(file_path)
        return self._parse_docx(file_path)

    def _parse_doc(self, file_path: Path) -> list[Document]:
        """解析 .doc 文件

        策略:
          1. 尝试 libreoffice 转换 (最佳质量)
          2. 降级到 olefile 纯 Python 提取 (无需系统依赖)
        """
        # 尝试 libreoffice
        try:
            return self._parse_doc_libreoffice(file_path)
        except (FileNotFoundError, OSError) as e:
            logger.info("LibreOffice not available, falling back to olefile: %s", e)

        # 降级到 olefile
        return self._parse_doc_olefile(file_path)

    def _parse_doc_libreoffice(self, file_path: Path) -> list[Document]:
        """用 LibreOffice 将 .doc 转为 .docx 后再解析"""
        import shutil

        if not shutil.which("libreoffice"):
            raise FileNotFoundError("libreoffice not found in PATH")

        with tempfile.TemporaryDirectory() as tmp_dir:
            result = subprocess.run(
                [
                    "libreoffice",
                    "--headless",
                    "--convert-to",
                    "docx",
                    "--outdir",
                    tmp_dir,
                    str(file_path),
                ],
                capture_output=True,
                text=True,
                timeout=60,
            )
            if result.returncode != 0:
                raise RuntimeError(f"LibreOffice conversion failed: {result.stderr}")

            docx_path = Path(tmp_dir) / f"{file_path.stem}.docx"
            if not docx_path.exists():
                raise FileNotFoundError(f"Converted file not found: {docx_path}")

            docs = self._parse_docx(docx_path)
            for doc in docs:
                doc.source = str(file_path)
            return docs

    def _parse_doc_olefile(self, file_path: Path) -> list[Document]:
        """用 olefile 从 .doc 的 OLE 容器中提取文本

        .doc 文件是 OLE2 格式，文本存储在 WordDocument stream 中。
        这里使用简化的二进制解析提取纯文本。
        """
        try:
            import olefile
        except ImportError:
            raise ImportError(
                "olefile not installed. Install with: pip install olefile"
            )

        source = str(file_path)

        if not olefile.isOleFile(str(file_path)):
            raise ValueError(f"Not a valid OLE file: {file_path}")

        ole = olefile.OleFileIO(str(file_path))
        text = ""

        try:
            # 方法 1: 尝试从 WordDocument stream 提取
            if ole.exists("WordDocument"):
                word_stream = ole.openstream("WordDocument").read()
                # 尝试提取可读文本（简化方法）
                text = self._extract_text_from_word_stream(word_stream, ole)

            # 方法 2: 如果上面没提取到，尝试所有文本流
            if not text.strip():
                text = self._extract_text_from_all_streams(ole)

        finally:
            ole.close()

        if not text.strip():
            return [
                Document(
                    source=source,
                    content=f"> (无法从 .doc 文件提取文本: {file_path.name})",
                    content_type=ContentType.PARAGRAPH,
                    metadata={"parse_method": "olefile", "partial": True},
                )
            ]

        # 按段落拆分
        documents = []
        paragraphs = re.split(r"\n{2,}", text.strip())
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            documents.append(
                Document(
                    source=source,
                    content=para,
                    content_type=ContentType.PARAGRAPH,
                    metadata={"parse_method": "olefile"},
                )
            )

        return documents

    def _extract_text_from_word_stream(self, word_stream: bytes, ole) -> str:
        """从 WordDocument stream 和对应的 Table stream 提取文本

        Word 二进制格式(FIB)中，文本位置由 FIB 的 ccpText 等字段决定。
        简化实现：提取 WordDocument 和 0Table/1Table 中的可读文本。
        """
        text_parts = []

        # 尝试读取 FIB 头部获取文本长度等信息
        try:
            # FIB base: magic number at offset 0 should be 0xA5EC
            if len(word_stream) >= 12:
                magic = struct.unpack_from("<H", word_stream, 0)[0]
                if magic == 0xA5EC:
                    # ccpText at offset in FIB varies by version
                    # 简化：直接尝试从 stream 提取 Unicode/ANSI 文本
                    pass
        except struct.error:
            pass

        # 提取 stream 中的可读文本
        # Word binary 文档中文本通常以 UTF-16LE 或 CP1252 编码
        try:
            # 尝试 UTF-16LE 解码（Word 97+ 通常用 Unicode）
            decoded = word_stream.decode("utf-16-le", errors="ignore")
            # 过滤控制字符，保留可读文本
            readable = []
            for ch in decoded:
                if ch in ("\n", "\r", "\t") or (ord(ch) >= 0x20 and ord(ch) < 0xFFFE):
                    readable.append(ch)
            text = "".join(readable)
            # 清理
            text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", text)
            text = re.sub(r"\r\n?", "\n", text)
            if len(text.strip()) > 50:
                text_parts.append(text.strip())
        except Exception:
            pass

        # 也尝试从 Table stream (0Table 或 1Table) 获取补充信息
        for table_name in ["0Table", "1Table"]:
            if ole.exists(table_name):
                try:
                    table_stream = ole.openstream(table_name).read()
                    decoded = table_stream.decode("utf-16-le", errors="ignore")
                    readable = "".join(
                        ch
                        for ch in decoded
                        if ch in ("\n", "\r", "\t") or (0x20 <= ord(ch) < 0xFFFE)
                    )
                    readable = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", readable)
                    if len(readable.strip()) > len("".join(text_parts)):
                        text_parts = [readable.strip()]
                except Exception:
                    pass

        return "\n\n".join(text_parts)

    def _extract_text_from_all_streams(self, ole) -> str:
        """遍历所有 OLE stream，尝试提取可读文本"""
        best_text = ""

        for stream_path in ole.listdir():
            try:
                data = ole.openstream(stream_path).read()
                # 尝试 UTF-16LE
                text = data.decode("utf-16-le", errors="ignore")
                text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", text)
                text = re.sub(r"\r\n?", "\n", text)
                # 过滤出有意义的文本（至少包含一些中文或可读字符）
                readable_chars = sum(
                    1 for c in text if c.isalnum() or c in "，。！？、；：" "''（）"
                )
                if readable_chars > len(best_text):
                    best_text = text.strip()
            except Exception:
                continue

        return best_text

    def _parse_docx(self, file_path: Path) -> list[Document]:
        """解析 .docx 文件，按段落和表格提取

        内嵌图片策略：
          - 先扫描统计文本量和图片数
          - 字多图少 → 跳过图片
          - 图多字少 → 合批调用 AI 描述图片
        """
        from .image_describer import (
            should_parse_images,
            describe_images_batch,
            ImageItem,
        )

        docx = DocxDocument(str(file_path))
        documents: list[Document] = []
        source = str(file_path)

        # --- 第一遍：扫描统计文本量和图片数 ---
        total_text = ""
        image_count = 0
        for element in docx.element.body:
            if not isinstance(element.tag, str):
                continue
            tag = element.tag.split("}")[-1]
            if tag == "p":
                from docx.text.paragraph import Paragraph

                para = Paragraph(element, docx)
                total_text += para.text
                # 检测段落内的 inline 图片
                blip_els = element.findall(
                    ".//a:blip",
                    {"a": "http://schemas.openxmlformats.org/drawingml/2006/main"},
                )
                image_count += len(blip_els)
            elif tag == "tbl":
                from docx.table import Table

                try:
                    tbl = Table(element, docx)
                    for row in tbl.rows:
                        for cell in row.cells:
                            total_text += cell.text
                except Exception as e:
                    logger.warning("跳过畸形表格 (第一遍扫描): %s", e)

        parse_images = should_parse_images(len(total_text.strip()), image_count)

        # --- 第二遍：提取文本 + 收集图片 ---
        image_items: list[ImageItem] = []
        image_slots: list[int] = []

        for element in docx.element.body:
            if not isinstance(element.tag, str):
                continue
            tag = element.tag.split("}")[-1]

            if tag == "p":
                doc = self._parse_paragraph(element, docx, source)
                if doc and doc.content.strip():
                    documents.append(doc)

                # 收集段落内嵌图片
                if parse_images:
                    blip_els = element.findall(
                        ".//a:blip",
                        {"a": "http://schemas.openxmlformats.org/drawingml/2006/main"},
                    )
                    for blip in blip_els:
                        rId = blip.get(
                            "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                        )
                        if rId and rId in docx.part.rels:
                            img_part = docx.part.rels[rId].target_part
                            image_items.append(
                                ImageItem(
                                    image_bytes=img_part.blob,
                                    mime_type=img_part.content_type or "image/png",
                                    context=f"DOCX文档{file_path.name}",
                                )
                            )
                            documents.append(None)  # type: ignore[arg-type]
                            image_slots.append(len(documents) - 1)

            elif tag == "tbl":
                try:
                    doc = self._parse_table(element, docx, source)
                    if doc:
                        documents.append(doc)
                except Exception as e:
                    logger.warning("跳过畸形表格: %s", e)

        # --- 合批描述图片 ---
        if image_items:
            describe_images_batch(image_items, self.config)
            for slot_idx, item in zip(image_slots, image_items):
                if item.description:
                    documents[slot_idx] = Document(
                        source=source,
                        content=item.description,
                        content_type=ContentType.IMAGE,
                        metadata={"embedded_image": True},
                    )
            documents = [d for d in documents if d is not None]

        return documents

    def _parse_paragraph(
        self, element, docx: DocxDocument, source: str
    ) -> Document | None:
        """解析段落元素"""
        from docx.text.paragraph import Paragraph

        para = Paragraph(element, docx)
        text = para.text.strip()
        if not text:
            return None

        # 判断标题级别
        style_name = (para.style.name or "").lower()
        if style_name.startswith("heading"):
            try:
                level = int(style_name.replace("heading", "").strip())
            except ValueError:
                level = 1
            content = f"{'#' * level} {text}"
            content_type = ContentType.HEADING
        elif style_name.startswith("list"):
            content = f"- {text}"
            content_type = ContentType.LIST
        else:
            content = text
            content_type = ContentType.PARAGRAPH

        return Document(
            source=source,
            content=content,
            content_type=content_type,
        )

    def _parse_table(self, element, docx: DocxDocument, source: str) -> Document | None:
        """解析表格元素，输出 Markdown 表格"""
        from docx.table import Table

        table = Table(element, docx)

        rows = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            rows.append(cells)

        if not rows:
            return None

        # 构建 Markdown 表格
        lines = []
        # 表头
        lines.append("| " + " | ".join(rows[0]) + " |")
        lines.append("| " + " | ".join(["---"] * len(rows[0])) + " |")
        # 数据行
        for row in rows[1:]:
            # 补齐列数
            while len(row) < len(rows[0]):
                row.append("")
            lines.append("| " + " | ".join(row[: len(rows[0])]) + " |")

        return Document(
            source=source,
            content="\n".join(lines),
            content_type=ContentType.TABLE,
        )
